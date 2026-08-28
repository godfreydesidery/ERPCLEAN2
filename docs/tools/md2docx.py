#!/usr/bin/env python
"""
md2docx.py — convert a (controlled-subset) Markdown file to a styled MS Word .docx.

Supports the subset the ERPCLEAN2 docs use: ATX headings (#..####), '- ' bullets,
'- [ ] ' task lists, '1. ' numbered lists, GitHub pipe tables, **bold**, *italic*,
`inline code`, [links](target), ``` fenced code (including inside a blockquote),
'> ' blockquotes, '---' horizontal rules, block images, and a leading cover (first H1)
+ a static, clickable Table of Contents (chapters + sections, hyperlinked to per-heading
bookmarks — always visible on open, no field update needed).

Soft-wrapped prose is joined: consecutive non-blank lines form ONE Word paragraph, the
way Markdown means them, rather than one paragraph per source line. Indented lines
following a list item continue that item.

Usage:
    python docs/tools/md2docx.py <input.md> <output.docx> ["Document Title"] ["Subtitle"]

Requires: python-docx (pip install python-docx).
"""
import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)$')
TABLE_SEP_RE = re.compile(r'^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$')
TASK_RE = re.compile(r'^(\s*)[-*]\s+\[([ xX])\]\s+(.*)$')
BULLET_RE = re.compile(r'^(\s*)[-*]\s+(.*)$')
NUM_RE = re.compile(r'^(\s*)\d+\.\s+(.*)$')
QUOTE_RE = re.compile(r'^\s*>\s?')
# A block-level Markdown image on its own line: ![alt text](relative/path.png)
IMAGE_RE = re.compile(r'^\s*!\[(.*?)\]\(([^)]+)\)\s*$')

# Inline tokens, in precedence order. Links come first so [**bold**](url) keeps its link;
# code before italic so `a*b` is never read as emphasis.
INLINE_RE = re.compile(
    r'(?P<link>\[[^\]\n]+\]\([^)\s]+\))'
    r'|(?P<bold>\*\*(?:[^*]|\*(?!\*))+\*\*)'
    r'|(?P<code>`[^`\n]+`)'
    # Italic: a single * hugging non-space text, not adjacent to a word character or
    # another *. Deliberately conservative — glob patterns like *.sql and feat/** live in
    # these documents and must not be eaten.
    r'|(?P<ital>(?<![\w*])\*(?![\s*])[^*\n]+?(?<![\s*])\*(?![\w*]))'
)
LINK_PARTS_RE = re.compile(r'^\[([^\]]+)\]\(([^)]+)\)$')

# Page text width on Letter with default 1-inch margins ≈ 6.5"; cap a touch under that.
IMAGE_MAX_WIDTH_IN = 6.2
# Table-of-contents depth: chapters (H1) + sections (H2). Kept in lock-step between the
# TOC list and the per-heading ordinal bookmarks so their anchors line up.
TOC_DEPTH = 2
LINK_COLOR = RGBColor(0x11, 0x55, 0xCC)
CODE_COLOR = RGBColor(0xB0, 0x30, 0x60)

# slug -> bookmark name, filled while rendering headings so [text](#anchor) can jump.
_SLUGS = {}


def slugify(text):
    """GitHub's heading-anchor slug: lowercase, drop punctuation, then replace EACH space
    with a hyphen. Runs of spaces are NOT collapsed — dropping a '·' from '1 · Overview'
    leaves two spaces and therefore '1--overview', which is exactly the anchor our
    Markdown cross-references use. Collapsing them here silently breaks every one."""
    s = re.sub(r'`|\*|_', '', text).strip().lower()
    s = re.sub(r'[^\w\s-]', '', s)
    return s.replace(' ', '-')


def collect_toc(lines, max_level=3):
    """Headings (levels 1..max_level) in document order, skipping fenced code blocks —
    the source for the static Table of Contents (and matched 1:1 by the ordinal bookmarks
    the body render adds to each heading, so the anchors line up)."""
    entries, in_code = [], False
    for line in lines:
        if _dequote(line).strip().startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            continue
        m = HEADING_RE.match(line)
        if m and len(m.group(1)) <= max_level:
            entries.append((len(m.group(1)), m.group(2).strip()))
    return entries


def _dequote(line):
    """Strip one leading '> ' blockquote marker."""
    return QUOTE_RE.sub('', line, count=1)


def add_bookmark(paragraph, name, bmid):
    """Wrap a heading paragraph in a Word bookmark so TOC hyperlinks can jump to it."""
    start = OxmlElement('w:bookmarkStart'); start.set(qn('w:id'), str(bmid)); start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), str(bmid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_entry(doc, text, anchor, level):
    """A clickable, indented TOC line that links to the heading's bookmark. Rendered
    statically at build time so the Table of Contents is ALWAYS visible and navigable on
    open — no 'right-click → Update Field' step."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28 * (level - 1))
    p.paragraph_format.space_after = Pt(2)
    hy = OxmlElement('w:hyperlink'); hy.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    if level == 1:
        rpr.append(OxmlElement('w:b'))
    color = OxmlElement('w:color'); color.set(qn('w:val'), '1155CC'); rpr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rpr.append(u)
    r.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); hy.append(r); p._p.append(hy)


def _style_run(run, bold=False, italic=False, code=False, link=False):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = CODE_COLOR
    if link:
        run.font.color.rgb = LINK_COLOR
        run.underline = True


def _add_plain(container, text, bold=False, italic=False, code=False, link=False):
    """Add a run to a paragraph OR to a <w:hyperlink> element."""
    if hasattr(container, 'add_run'):          # a real paragraph
        run = container.add_run(text)
        _style_run(run, bold, italic, code, link)
        return
    # inside a hyperlink element: build the run by hand
    r = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    if bold:
        rpr.append(OxmlElement('w:b'))
    if italic:
        rpr.append(OxmlElement('w:i'))
    if code:
        rf = OxmlElement('w:rFonts')
        for a in ('w:ascii', 'w:hAnsi'):
            rf.set(qn(a), 'Consolas')
        rpr.append(rf)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '1155CC'); rpr.append(col)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rpr.append(u)
    r.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t)
    container.append(r)


def add_link(paragraph, text, target):
    """Render [text](target) as a real Word hyperlink.

    '#anchor' becomes an internal jump when that heading exists in this document; a URL or
    a relative path becomes an external hyperlink (a relative path resolves against the
    .docx's own folder, so 'ops/backup-restore.md' works when the document sits in docs/).
    An anchor with no matching heading degrades to plain styled text rather than a dead
    link.
    """
    hy = OxmlElement('w:hyperlink')
    if target.startswith('#'):
        bookmark = _SLUGS.get(target[1:].lower())
        if not bookmark:
            _add_inline(paragraph, text, base_link=True)      # nothing to point at
            return
        hy.set(qn('w:anchor'), bookmark)
    else:
        r_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
        hy.set(qn('r:id'), r_id)
    _add_inline(hy, text, base_link=True)
    paragraph._p.append(hy)


def _add_inline(container, text, base_bold=False, base_italic=False, base_link=False):
    """Render inline Markdown (links, **bold**, *italic*, `code`) into a paragraph or a
    hyperlink element. Recurses so **bold with `code` inside** keeps both."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_plain(container, text[pos:m.start()], base_bold, base_italic, False, base_link)
        kind = m.lastgroup
        tok = m.group(0)
        if kind == 'link':
            lm = LINK_PARTS_RE.match(tok)
            if hasattr(container, 'add_run'):
                add_link(container, lm.group(1), lm.group(2))
            else:                                   # a link inside a link: keep the text
                _add_inline(container, lm.group(1), base_bold, base_italic, True)
        elif kind == 'bold':
            _add_inline(container, tok[2:-2], True, base_italic, base_link)
        elif kind == 'ital':
            _add_inline(container, tok[1:-1], base_bold, True, base_link)
        else:                                        # code
            _add_plain(container, tok[1:-1], base_bold, base_italic, True, base_link)
        pos = m.end()
    if pos < len(text):
        _add_plain(container, text[pos:], base_bold, base_italic, False, base_link)


def add_runs(paragraph, text):
    """Render one line of inline Markdown into an existing paragraph."""
    _add_inline(paragraph, text)


def add_image(doc, alt, src, base):
    """Embed a block-level Markdown image, centered, with the alt text as a caption.

    Paths resolve relative to `base` (set by the build via MD2DOCX_IMG_BASE). A missing or
    unreadable file degrades to an italic placeholder rather than aborting the build.
    """
    path = src if os.path.isabs(src) else os.path.normpath(os.path.join(base, src))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        try:
            p.add_run().add_picture(path, width=Inches(IMAGE_MAX_WIDTH_IN))
        except Exception as exc:  # corrupt/unsupported image — don't fail the whole doc
            r = p.add_run(f'[image unavailable: {alt or src} — {exc}]'); r.italic = True
    else:
        r = p.add_run(f'[missing image: {src}]'); r.italic = True
        sys.stderr.write(f'WARN md2docx: missing image {path}\n')
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(alt)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def flush_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=cols)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ''
            para = cells[ci].paragraphs[0]
            add_runs(para, cell_text.strip())
            if ri == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def add_code_block(doc, code_lines, indent=0.0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run('\n'.join(code_lines))
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)


def add_quote_shading(paragraph):
    """A light band behind a blockquote so callouts read as callouts."""
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'FBF6E7')
    paragraph._p.get_or_add_pPr().append(shd)


def convert(md_path, docx_path, title=None, subtitle=None):
    # Image base: where chapter-relative `images/...` links are anchored.
    img_base = os.environ.get('MD2DOCX_IMG_BASE') or os.path.dirname(os.path.abspath(md_path))

    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    _SLUGS.clear()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ── Cover page + static, clickable Table of Contents ──
    if title:
        for _ in range(6):
            doc.add_paragraph()
        h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run(title); r.bold = True; r.font.size = Pt(30)
        if subtitle:
            s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = s.add_run(subtitle); sr.font.size = Pt(15); sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_page_break()
        doc.add_heading('Contents', level=1)
        for idx, (lvl, text) in enumerate(collect_toc(lines, max_level=TOC_DEPTH)):
            add_toc_entry(doc, re.sub(r'`|\*\*', '', text), f'_Toc{idx:04d}', lvl)
        doc.add_page_break()

    # ── Pass 1: every heading's slug -> a stable bookmark name, so an inline
    #    [text](#anchor) written anywhere in the document can resolve to it.
    in_code = False
    for line in lines:
        if _dequote(line).strip().startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            continue
        m = HEADING_RE.match(line)
        if m:
            _SLUGS.setdefault(slugify(m.group(2).strip()), f'_md{len(_SLUGS):04d}')

    # ── Pass 2: render ──
    # Buffered paragraph state. Markdown soft-wraps: consecutive non-blank lines are ONE
    # paragraph. Without this every wrapped source line became its own Word paragraph.
    buf, buf_kind, buf_indent = [], None, 0

    def flush_para():
        nonlocal buf, buf_kind, buf_indent
        if not buf:
            buf_kind = None
            return
        text = ' '.join(part.strip() for part in buf if part.strip())
        buf = []
        kind, indent = buf_kind, buf_indent
        buf_kind, buf_indent = None, 0
        if not text:
            return
        if kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5 + 0.25 * (indent // 2 - 1))
        elif kind == 'number':
            p = doc.add_paragraph(style='List Number')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5 + 0.25 * (indent // 2 - 1))
        elif kind == 'task':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            box = p.add_run(text[:1] + '  ')          # the ☐ / ☑ glyph carried in text[0]
            box.font.name = 'Segoe UI Symbol'
            text = text[1:].lstrip()
        elif kind == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.15)
            add_quote_shading(p)
        else:
            p = doc.add_paragraph()
        add_runs(p, text)

    i = 0
    toc_idx = 0            # ordinal of the next H1/H2 — must track collect_toc()'s order
    table_buf = []
    in_code = False
    code_buf = []
    code_indent = 0.0
    code_quoted = False

    while i < len(lines):
        line = lines[i]
        deq = _dequote(line) if QUOTE_RE.match(line) else line

        # ── fenced code (also when the fence sits inside a blockquote) ──
        if deq.strip().startswith('```'):
            if in_code:
                add_code_block(doc, code_buf, code_indent)
                code_buf = []; in_code = False; code_indent = 0.0; code_quoted = False
            else:
                flush_para()
                if table_buf:
                    flush_table(doc, table_buf); table_buf = []
                in_code = True
                code_quoted = QUOTE_RE.match(line) is not None
                lead = len(deq) - len(deq.lstrip())
                code_indent = 0.4 if lead >= 2 else 0.0
            i += 1
            continue
        if in_code:
            code_buf.append(_dequote(line) if code_quoted else line)
            i += 1
            continue

        stripped = line.strip()

        # ── pipe tables ──
        if stripped.startswith('|') and '|' in stripped[1:]:
            flush_para()
            if TABLE_SEP_RE.match(line):
                i += 1; continue                       # skip the |---|---| separator
            table_buf.append([c for c in stripped.strip('|').split('|')])
            i += 1; continue
        elif table_buf:
            flush_table(doc, table_buf); table_buf = []

        # ── headings ──
        m = HEADING_RE.match(line)
        if m:
            flush_para()
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading('', level=min(level, 4))
            add_runs(h, text)
            slug_bm = _SLUGS.get(slugify(text))
            if slug_bm:
                add_bookmark(h, slug_bm, 5000 + len(_SLUGS) + i)
            if level <= TOC_DEPTH:                     # ordinal bookmark for the TOC list
                add_bookmark(h, f'_Toc{toc_idx:04d}', toc_idx)
                toc_idx += 1
            i += 1; continue

        # ── block image ──
        im = IMAGE_RE.match(line)
        if im:
            flush_para()
            add_image(doc, im.group(1).strip(), im.group(2).strip(), img_base)
            i += 1; continue

        # ── horizontal rule ──
        if stripped in ('---', '***', '___'):
            flush_para()
            doc.add_paragraph().add_run().add_break()
            i += 1; continue

        # ── blank line ends whatever was being built ──
        if stripped == '':
            flush_para()
            i += 1; continue

        # ── blockquote ──
        if QUOTE_RE.match(line):
            body = _dequote(line)
            if body.strip() == '':
                flush_para(); i += 1; continue
            if buf_kind != 'quote':
                flush_para(); buf_kind = 'quote'
            buf.append(body)
            i += 1; continue

        # ── list items ──
        tm = TASK_RE.match(line)
        if tm:
            flush_para()
            buf_kind, buf_indent = 'task', len(tm.group(1))
            buf.append(('☑' if tm.group(2).lower() == 'x' else '☐') + ' ' + tm.group(3))
            i += 1; continue

        bm = BULLET_RE.match(line)
        if bm:
            flush_para()
            buf_kind, buf_indent = 'bullet', len(bm.group(1))
            buf.append(bm.group(2))
            i += 1; continue

        nm = NUM_RE.match(line)
        if nm:
            flush_para()
            buf_kind, buf_indent = 'number', len(nm.group(1))
            buf.append(nm.group(2))
            i += 1; continue

        # ── plain text: continues the current block (soft wrap), or starts one ──
        indented = len(line) - len(line.lstrip()) >= 2
        if buf_kind is None:
            buf_kind = 'p'
        elif buf_kind in ('bullet', 'number', 'task', 'quote') and not indented:
            flush_para(); buf_kind = 'p'               # a new paragraph, not a continuation
        buf.append(line)
        i += 1

    flush_para()
    if table_buf:
        flush_table(doc, table_buf)

    doc.save(docx_path)
    print(f'Wrote {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(2)
    convert(sys.argv[1], sys.argv[2],
            sys.argv[3] if len(sys.argv) > 3 else None,
            sys.argv[4] if len(sys.argv) > 4 else None)
