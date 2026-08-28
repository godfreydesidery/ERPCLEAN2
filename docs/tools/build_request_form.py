# -*- coding: utf-8 -*-
"""Builds a plain-language Word form the customer fills in to report a
problem or ask for something new. No technical jargon anywhere.

Writing room is the priority: answer boxes are sized for real handwriting
(one ruled line = LINE_H), and page count is allowed to grow to suit.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHK = u"\u2610"          # empty checkbox
ARROW = u"\u2192"
NDASH = u"\u2013"

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BAR_FILL = "1F3864"
LABEL_FILL = "EDF0F7"
NOTE_FILL = "FFF8E1"

GRID = 'Table Grid'

LABEL_W = Cm(5.0)
VALUE_W = Cm(12.0)
FULL_W = Cm(17.0)

LINE_PT = 26             # height of one writing line, in points
LINE_H = LINE_PT / 28.35  # ...the same thing in centimetres


# ---------- low-level helpers ----------

def shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=10.5, color=None, italic=False,
              space_after=2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    first = True
    for line in text.split("\n"):
        if not first:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(space_after)
        first = False
        run = p.add_run(line)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
    return p


def writing_lines(cell, n):
    """Give an answer cell n generously spaced lines to write on."""
    cell.text = ""
    for i in range(n):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(LINE_PT)


def set_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w


def min_height(row, cm):
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Cm(cm)


def para(doc, text="", bold=False, size=10.5, italic=False, color=None,
         space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
    return p


def section_bar(doc, text):
    t = doc.add_table(rows=1, cols=1)
    set_widths(t, [FULL_W])
    c = t.rows[0].cells[0]
    shade(c, BAR_FILL)
    cell_text(c, text, bold=True, size=11, color=WHITE)
    para(doc, space_after=2)
    return t


def form_table(doc, rows):
    """rows: list of (label, kind, payload)
       kind 'lines' -> payload = number of writing lines
       kind 'text'  -> payload = pre-printed text (e.g. tick boxes)"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = GRID
    set_widths(t, [LABEL_W, VALUE_W])
    for i, (label, kind, payload) in enumerate(rows):
        lc, vc = t.rows[i].cells
        shade(lc, LABEL_FILL)
        cell_text(lc, label, bold=True, size=10)
        if kind == 'lines':
            writing_lines(vc, payload)
            min_height(t.rows[i], LINE_H * payload)
        else:
            cell_text(vc, payload, size=10)
            min_height(t.rows[i], 0.85)
    para(doc, space_after=6)
    return t


def note_box(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = GRID
    set_widths(t, [FULL_W])
    c = t.rows[0].cells[0]
    shade(c, NOTE_FILL)
    cell_text(c, text, size=9.5, italic=True)
    para(doc, space_after=6)
    return t


def hairline(doc):
    """Word insists on a paragraph after a table. On a near-full page a
    normal-sized one spills over and costs a blank page, so keep it tiny."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run("").font.size = Pt(1)
    return p


def open_box(doc, height_cm, hint=None, gap_pt=6):
    """A big empty bordered box - for a screenshot or free writing."""
    t = doc.add_table(rows=1, cols=1)
    t.style = GRID
    set_widths(t, [FULL_W])
    c = t.rows[0].cells[0]
    if hint:
        cell_text(c, hint, size=9.5, italic=True, color=GREY)
    else:
        writing_lines(c, int(height_cm / LINE_H))
    min_height(t.rows[0], height_cm)
    if gap_pt:
        para(doc, space_after=gap_pt)
    else:
        hairline(doc)
    return t


def add_page_number_footer(section, left_text):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(left_text + "     |     page ")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    p._p.append(fld)


# ---------- document ----------

doc = Document()

st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

sec = doc.sections[0]
sec.top_margin = Cm(1.7)
sec.bottom_margin = Cm(1.5)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.0)
add_page_number_footer(sec, "Report a problem or request a change")

# ---- Page 1: cover + how to use ----

p = para(doc, space_after=0)
r = p.add_run("[ Your company name ]")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

p = para(doc, space_after=2, space_before=4)
r = p.add_run("Report a Problem")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY
r2 = p.add_run("  or  ")
r2.font.size = Pt(16)
r2.font.color.rgb = GREY
r3 = p.add_run("Request a Change")
r3.bold = True
r3.font.size = Pt(26)
r3.font.color.rgb = NAVY

para(doc,
     "Use this form to tell us about anything in the system that is not "
     "working the way you expect, or anything new you would like it to do.",
     size=11.5, color=GREY, space_after=14)

section_bar(doc, "How to use this form")

steps = [
    ("1.", "Use one form for each separate item. If you have three things to "
           "tell us, fill in three forms."),
    ("2.", "Write as much as you like. If you are typing, the boxes grow as "
           "you type " + ARROW + " they will never run out of room. If you "
           "are writing by hand and need more space, use the extra pages at "
           "the back."),
    ("3.", "Fill in as much as you can. If you do not know an answer, leave "
           "it blank " + ARROW + " a half-filled form is much better than no "
           "form at all."),
    ("4.", "Copy the blank form (every page from the one headed \"Form\" to "
           "the end) as many times as you need. Select those pages, press "
           "Ctrl+C, put the cursor at the very end of the document and press "
           "Ctrl+V."),
    ("5.", "Add a picture of the screen if you can. Hold the Windows key + "
           "Shift and press S, drag a box around what you can see, then click "
           "inside the picture box on the form and press Ctrl+V."),
    ("6.", "Write the item on the summary list on the next page too, so we "
           "both know what is still outstanding."),
    ("7.", "Save the file and send it to: ______________________________ "
           "(email / WhatsApp)."),
]
for num, txt in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.first_line_indent = Cm(-0.9)
    rn = p.add_run(num + "  ")
    rn.bold = True
    rn.font.color.rgb = NAVY
    p.add_run(txt)

para(doc, space_after=6)
section_bar(doc, "How urgent is it? " + NDASH + " what we mean")

urg = [
    ("How urgent", "What it means"),
    ("Stopping work", "We cannot carry on. Sales, deliveries or payments are "
                      "on hold right now."),
    ("Slowing us down", "We can still work, but it takes much longer or we "
                        "have to do the same thing twice."),
    ("Annoying", "It works, but it is awkward, confusing or untidy."),
    ("Would be nice", "Not urgent. Something that would help us when there "
                      "is time for it."),
]
t = doc.add_table(rows=len(urg), cols=2)
t.style = GRID
set_widths(t, [Cm(4.2), Cm(12.8)])
for i, (a, b) in enumerate(urg):
    ca, cb = t.rows[i].cells
    if i == 0:
        shade(ca, BAR_FILL)
        shade(cb, BAR_FILL)
        cell_text(ca, a, bold=True, size=10, color=WHITE)
        cell_text(cb, b, bold=True, size=10, color=WHITE)
    else:
        shade(ca, LABEL_FILL)
        cell_text(ca, a, bold=True, size=10)
        cell_text(cb, b, size=10)
    min_height(t.rows[i], 0.6)

para(doc, space_after=8)
note_box(doc,
         "Please do not write passwords, PIN numbers or bank details on this "
         "form. If we ever need anything like that, we will ask you for it "
         "separately.")

# ---- Page 2: running summary list ----

doc.add_page_break()

para(doc, "Summary list", bold=True, size=18, color=NAVY, space_after=2)
para(doc,
     "Write one line here for every form you send us. We will fill in the "
     "last column and send the list back, so you can always see where each "
     "item stands.",
     size=10.5, color=GREY, space_after=10)

cols = ["No.", "Date sent", "What it is about (a few words)",
        "Problem or\nrequest?", "How urgent", "Where it stands\n(we fill in)"]
widths = [Cm(1.0), Cm(2.1), Cm(6.1), Cm(2.2), Cm(2.5), Cm(3.1)]
t = doc.add_table(rows=16, cols=6)
t.style = GRID
set_widths(t, widths)
for j, h in enumerate(cols):
    c = t.rows[0].cells[j]
    shade(c, BAR_FILL)
    cell_text(c, h, bold=True, size=9.5, color=WHITE)
min_height(t.rows[0], 0.9)
for i in range(1, 16):
    min_height(t.rows[i], 1.15)
    cell_text(t.rows[i].cells[0], str(i), size=9.5, color=GREY)

# ---- The form itself ----

doc.add_page_break()

p = para(doc, space_after=2)
r = p.add_run("Form")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY
r = p.add_run("      No. ______      Date: ______________________")
r.font.size = Pt(11)
r.font.color.rgb = GREY

para(doc,
     "One item per form, please. Fill in Part 3 if something is not working. "
     "Fill in Part 4 if you are asking for something new or different. You do "
     "not need to fill in both.",
     size=10, italic=True, color=GREY, space_after=10)

section_bar(doc, "Part 1  " + NDASH + "  Who is reporting this")
form_table(doc, [
    ("Your name", 'lines', 1),
    ("Your job / what you do", 'lines', 1),
    ("Branch, shop or office", 'lines', 1),
    ("Phone or email\n(in case we need to ask you more)", 'lines', 1),
])

section_bar(doc, "Part 2  " + NDASH + "  What is this about?")
form_table(doc, [
    ("Describe it in one line", 'lines', 2),
    ("Is this...",
     'text', CHK + "  Something is not working        " +
             CHK + "  Something new I would like\n" +
             CHK + "  A change to how it works now    " +
             CHK + "  A question " + NDASH + " I am not sure"),
    ("Which part of the system?",
     'text', CHK + "  Selling, invoices, customers      " +
             CHK + "  Buying, suppliers, orders\n" +
             CHK + "  Stock and the store               " +
             CHK + "  Money in and money out\n" +
             CHK + "  Reports and figures               " +
             CHK + "  The till / point of sale\n" +
             CHK + "  Staff, users and what they are allowed to see\n" +
             CHK + "  Other (please say): ______________________________"),
    ("How urgent is it?",
     'text', CHK + "  Stopping work    " + CHK + "  Slowing us down    " +
             CHK + "  Annoying    " + CHK + "  Would be nice"),
    ("How many people does it affect?",
     'text', CHK + "  Just me    " + CHK + "  A few of us    " +
             CHK + "  My whole branch    " + CHK + "  Everybody"),
])

doc.add_page_break()

section_bar(doc, "Part 3  " + NDASH + "  If something is not working")
note_box(doc,
         "The more of this you can fill in, the faster we can find it. "
         "Please write what you actually saw on the screen, not what you "
         "think caused it.")
form_table(doc, [
    ("What were you doing?\nWrite it as steps:\nfirst I..., then I...",
     'lines', 7),
    ("What happened?", 'lines', 5),
    ("What did you expect to happen instead?", 'lines', 5),
])

doc.add_page_break()

section_bar(doc, "Part 3  " + NDASH + "  If something is not working "
                 "(continued)")
form_table(doc, [
    ("Was there a message on the screen?\nPlease copy it word for word.",
     'lines', 4),
    ("When did you first notice it?", 'lines', 2),
    ("Does it happen every time?",
     'text', CHK + "  Every time    " + CHK + "  Most times    " +
             CHK + "  Now and again    " + CHK + "  It happened once"),
    ("Does it happen to other people, on other computers, or at other "
     "branches?",
     'text', CHK + "  Yes    " + CHK + "  No    " + CHK + "  I do not know"),
    ("Anything that helps us find it\ne.g. invoice number, item name or code, "
     "customer name, and the date and time it happened",
     'lines', 6),
])

doc.add_page_break()

section_bar(doc, "Part 4  " + NDASH + "  If you are asking for something new "
                 "or different")
form_table(doc, [
    ("What would you like to be able to do?", 'lines', 7),
    ("Why do you need it?\nWhat goes wrong today without it?", 'lines', 6),
    ("How do you manage at the moment?\ne.g. on paper, in Excel, by phone, "
     "or we just do without",
     'lines', 5),
    ("How often would it be used?",
     'text', CHK + "  Many times a day    " + CHK + "  Daily    " +
             CHK + "  Weekly    " + CHK + "  Monthly    " +
             CHK + "  Now and again"),
])

doc.add_page_break()

section_bar(doc, "Part 4  " + NDASH + "  If you are asking for something new "
                 "(continued)")
form_table(doc, [
    ("Is there a date you need it by?\nIf yes, please say why (audit, tax "
     "return, busy season, a new shop opening).",
     'lines', 4),
])

section_bar(doc, "Part 5  " + NDASH + "  Picture of the screen (if you have "
                 "one)")
open_box(doc, 10.5, hint="Click here and press Ctrl+V to paste your picture.")

doc.add_page_break()

section_bar(doc, "Part 6  " + NDASH + "  Anything else you want us to know")
open_box(doc, 12.0)

para(doc, space_after=4)
para(doc, "For our team " + NDASH + " please leave this part blank",
     bold=True, size=9.5, color=GREY, space_after=3)
t = doc.add_table(rows=2, cols=4)
t.style = GRID
set_widths(t, [Cm(4.25)] * 4)
heads = ["Received on", "Our reference", "Agreed action", "Expected by"]
for j, h in enumerate(heads):
    c = t.rows[0].cells[j]
    shade(c, LABEL_FILL)
    cell_text(c, h, bold=True, size=9)
min_height(t.rows[0], 0.55)
min_height(t.rows[1], 1.4)

# ---- Extra space pages ----

for _ in range(2):
    doc.add_page_break()
    p = para(doc, space_after=2)
    r = p.add_run("More space")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY
    r = p.add_run("      Form No. ______      Carrying on from Part ______")
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    para(doc,
         "Use this page if you ran out of room anywhere above. Please write "
         "which part you are continuing.",
         size=10, italic=True, color=GREY, space_after=8)
    open_box(doc, 20.5, gap_pt=0)

out = r"d:\My_Works\ERP\ERPCLEAN2\docs\templates\Report-a-Problem-or-Request-a-Change.docx"
doc.save(out)
print("saved:", out)
